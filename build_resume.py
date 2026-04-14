"""
Generate Vamshi_Yadav_Golla_Resume.docx matching the portfolio's resume.html
Uses python-docx. Run: python build_resume.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "Vamshi_Yadav_Golla_Resume.docx"

# Colors
INK = RGBColor(0x0F, 0x17, 0x2A)
INK2 = RGBColor(0x33, 0x41, 0x55)
MUTED = RGBColor(0x64, 0x74, 0x8B)
ACCENT = RGBColor(0x63, 0x66, 0xF1)

doc = Document()

# Page margins (A4, narrow margins for density)
for section in doc.sections:
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

# Set default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.font.color.rgb = INK2


def set_cell_border(cell, **kwargs):
    """Apply borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        border = OxmlElement(f"w:{edge}")
        for k, v in val.items():
            border.set(qn(f"w:{k}"), v)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading_rule(text):
    """Section heading with accent color and bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT
    run.font.name = "Calibri"
    # Letter spacing via xml
    rPr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "30")
    rPr.append(spacing)
    # Bottom border on paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E2E8F0")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_bullet(text_segments):
    """text_segments: list of (text, bold) tuples to build bold-inline bullets."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    for text, bold in text_segments:
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
        run.font.color.rgb = INK2
        if bold:
            run.bold = True
            run.font.color.rgb = INK


def add_entry_head(role, org_info, date):
    """Role on left (bold), date on right, then org line underneath."""
    # Create a table for the role / date row (cleaner than tab stops for alignment)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(13.5)
    table.columns[1].width = Cm(4)
    row = table.rows[0]
    # Left cell: role
    left = row.cells[0]
    left.width = Cm(13.5)
    lp = left.paragraphs[0]
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after = Pt(0)
    rrun = lp.add_run(role)
    rrun.bold = True
    rrun.font.size = Pt(10.5)
    rrun.font.color.rgb = INK
    # Right cell: date (right-aligned)
    right = row.cells[1]
    right.width = Cm(4)
    rp = right.paragraphs[0]
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after = Pt(0)
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    drun = rp.add_run(date)
    drun.font.size = Pt(9)
    drun.font.color.rgb = MUTED
    # Remove table borders
    for cell in row.cells:
        set_cell_border(
            cell,
            top={"val": "nil"},
            bottom={"val": "nil"},
            left={"val": "nil"},
            right={"val": "nil"},
        )
    # Organization line
    op = doc.add_paragraph()
    op.paragraph_format.space_before = Pt(0)
    op.paragraph_format.space_after = Pt(2)
    orun = op.add_run(org_info)
    orun.font.size = Pt(9.5)
    orun.font.color.rgb = INK2


def add_tech_line(stack, link=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    label = p.add_run("Stack: ")
    label.bold = True
    label.font.size = Pt(8.5)
    label.font.color.rgb = INK
    body = p.add_run(stack)
    body.font.size = Pt(8.5)
    body.font.color.rgb = MUTED
    if link:
        sep = p.add_run("  |  ")
        sep.font.size = Pt(8.5)
        sep.font.color.rgb = MUTED
        lnk = p.add_run(link)
        lnk.font.size = Pt(8.5)
        lnk.font.color.rgb = ACCENT


# ================== HEADER (name + contact grid) ==================
# Name
h = doc.add_paragraph()
h.paragraph_format.space_before = Pt(0)
h.paragraph_format.space_after = Pt(0)
name_run = h.add_run("Vamshi Yadav Golla")
name_run.bold = True
name_run.font.size = Pt(24)
name_run.font.color.rgb = INK
name_run.font.name = "Calibri"

# Title
t = doc.add_paragraph()
t.paragraph_format.space_before = Pt(0)
t.paragraph_format.space_after = Pt(4)
title_run = t.add_run("AI Agent Engineer & Automation Architect")
title_run.bold = True
title_run.font.size = Pt(11)
title_run.font.color.rgb = ACCENT

# Contact line
c = doc.add_paragraph()
c.paragraph_format.space_before = Pt(0)
c.paragraph_format.space_after = Pt(4)
contact_parts = [
    "London, UK",
    "Vamshiyadav2783@gmail.com",
    "+44 7887 132 409",
    "linkedin.com/in/vamshi-yadav869",
    "github.com/Vamshi-27072001",
    "vamshi-27072001.github.io/portfolio",
]
for i, part in enumerate(contact_parts):
    if i > 0:
        sep = c.add_run("  •  ")
        sep.font.size = Pt(9)
        sep.font.color.rgb = MUTED
    run = c.add_run(part)
    run.font.size = Pt(9)
    run.font.color.rgb = INK2

# Bottom rule under header
ruler = doc.add_paragraph()
ruler.paragraph_format.space_before = Pt(0)
ruler.paragraph_format.space_after = Pt(0)
pPr = ruler._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "12")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "0F172A")
pBdr.append(bottom)
pPr.append(pBdr)


# ================== SUMMARY ==================
add_heading_rule("Professional Summary")
s = doc.add_paragraph()
s.paragraph_format.space_after = Pt(2)
summary_segments = [
    ("AI Agent Engineer", True),
    (" specialising in autonomous AI systems built on ", False),
    ("n8n, LangChain, and GPT-4o", True),
    (". Shipped ", False),
    ("6 production projects", True),
    (" delivering ", False),
    ("£36K+/year in business impact", True),
    (", including a 99.98% cost-reduction chatbot, a zero-touch LinkedIn publishing pipeline, and an ML GPS-spoofing detector with ", False),
    ("99.95% accuracy", True),
    (". MSc Computer Science (UEL). Actively interviewing for UK-based AI Engineer, Automation Engineer, and ML Engineer roles — Graduate Visa eligible.", False),
]
for text, bold in summary_segments:
    r = s.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = INK2
    if bold:
        r.bold = True
        r.font.color.rgb = INK


# ================== EXPERIENCE ==================
add_heading_rule("Professional Experience")
add_entry_head(
    "Cyber Security Intern",
    "Lexnis Services Limited — Hertfordshire, UK",
    "Nov 2025 – Present",
)
for bullet in [
    [("Conducted market analysis of 10 UK vehicle tracking companies", True), (" across 5 benchmarking criteria, producing a decision-ready competitive matrix for leadership.", False)],
    [("Mapped 15+ product features per provider", True), (" and recommended best-fit GPS tracking solutions aligned to target customer segments.", False)],
    [("Identified a market gap supporting a £9,100/month operational model", True), (" — insight directly fed into the company's go-to-market strategy.", False)],
    [("Leveraged AI tools (GPT-4, Google suite) to produce research reports", True), (" that informed 3 major business decisions, accelerating analysis throughput 4×.", False)],
]:
    add_bullet(bullet)


# ================== PROJECTS ==================
add_heading_rule("Flagship Projects")

# --- Project 1 ---
add_entry_head(
    "NSP Cases — AI Email Enquiry Handler",
    "Autonomous customer-service pipeline  •  Production",
    "Apr 2026",
)
for bullet in [
    [("Engineered a 19-node n8n workflow that replaces a human customer-service agent", True), (" — reads Gmail enquiries, queries a 7-tab knowledge base, sends HTML replies end-to-end in ", False), ("under 60 seconds, zero human intervention", True), (".", False)],
    [("Integrated GPT-4o Vision for multimodal enquiry handling", True), (" — auto-analyses product drawings/images, removing the #1 friction point in technical sales enquiries.", False)],
    [("Built persona-driven AI agent (GPT-4.1-mini)", True), (" producing replies indistinguishable from a human expert; full Supabase PostgreSQL audit trail for regulatory traceability.", False)],
]:
    add_bullet(bullet)
add_tech_line(
    "n8n · LangChain · GPT-4.1-mini · GPT-4o Vision · Gmail OAuth2 · Google Sheets API · Supabase PostgreSQL",
    "github.com/Vamshi27072001/NSP-email-enquiry",
)

# --- Project 2 ---
add_entry_head(
    "Restaurant AI Agent — Taco Bell UK",
    "Zero-hallucination AI chatbot  •  Production",
    "Apr 2026",
)
for bullet in [
    [("Deployed an 11-node LangChain AI agent on Telegram (GPT-4o-mini)", True), (" handling menu, allergen, nutrition, and ingredient questions ", False), ("24/7 at sub-2-second latency", True), (".", False)],
    [("Drove operating costs from £12,000/year → £36/year (99.98% reduction)", True), (" while lifting coverage from business hours to always-on.", False)],
    [("Engineered strict tool-first routing for zero-hallucination accuracy", True), (" across 100+ menu items; full EU/UK allergen compliance with mandatory cross-contamination notices.", False)],
    [("Added persistent per-user memory via PostgreSQL", True), (" (100-message context window) for stateful conversations across sessions.", False)],
]:
    add_bullet(bullet)
add_tech_line(
    "n8n · LangChain · GPT-4o-mini · Telegram Bot API · Google Sheets/Docs · PostgreSQL",
    "github.com/Vamshi-27072001/restaurant-ai-agent",
)

# --- Project 3 ---
add_entry_head(
    "LinkedIn Post Manager — AI Content Engine",
    "Autonomous Telegram-to-LinkedIn publishing  •  Production",
    "Apr 2026",
)
for bullet in [
    [("Built a fully autonomous Telegram-to-LinkedIn publishing engine", True), (" (13-node n8n + GPT-4.1-mini + OpenAI Images API + LinkedIn ugcPosts API) eliminating every human touchpoint.", False)],
    [("Cut content creation time by 98% (25 min → 30 sec)", True), (" and saved ", False), ("£24,000/year in outsourced social media costs", True), (" at just £0.01 per post.", False)],
    [("Orchestrated a 5-API cross-platform media pipeline", True), (" with binary MIME-type correction, OpenAI image enhancement, and LinkedIn Assets API upload.", False)],
    [("PostgreSQL-backed LangChain memory (100-message window)", True), (" delivering 100% context retention; enforced production post structure (hook + 3 paragraphs + CTA + 10 hashtags) on every output.", False)],
]:
    add_bullet(bullet)
add_tech_line(
    "n8n · LangChain · GPT-4.1-mini · OpenAI Images API · LinkedIn REST API · Telegram · PostgreSQL · OAuth2",
    "github.com/Vamshi-27072001/telegram-linkedin-agent",
)

# --- Project 4 ---
add_entry_head(
    "ML-Based GPS Spoofing Detection for Drone Delivery Systems",
    "MSc Dissertation  •  University of East London",
    "Jun 2025 – Sep 2025",
)
for bullet in [
    [("Architected a production-grade ML security system for autonomous drones", True), (" — end-to-end pipeline from raw telemetry ingestion to live edge inference on Raspberry Pi.", False)],
    [("Engineered predictive features from 62,000+ live telemetry records", True), ("; isolated top 4 attack signals (GPS HDOP, yaw rate, throttle %, roll) via correlation and feature-importance analysis.", False)],
    [("Trained Random Forest & XGBoost classifiers to 99.95% / 99.90% accuracy", True), (" with false-positive rate <0.05% — exceeding aviation safety thresholds.", False)],
    [("Deployed real-time detection pipeline", True), (" with autonomous countermeasures (hover, reroute, return-to-base); validated against simulated spoofing attacks.", False)],
]:
    add_bullet(bullet)
add_tech_line("Python · Scikit-learn · XGBoost · Random Forest · Pandas · NumPy · Raspberry Pi · Matplotlib/Seaborn")

# --- Project 5 ---
add_entry_head(
    "Disaster Tweet Classification — NLP",
    "Kaggle competition",
    "2025",
)
for bullet in [
    [("Built a production-ready NLP classifier for real-time disaster detection", True), (" — 87% accuracy (+15% over baseline) on 7,000+ tweets.", False)],
    [("Engineered a 6-stage text-preprocessing pipeline", True), (" (tokenization, stopword removal, stemming, lemmatization, regex cleaning, TF-IDF), materially lifting precision/recall.", False)],
    [("Recovered 90% of missing geolocation data", True), (" via custom imputation heuristics, unlocking geospatial analysis for emergency-response use cases.", False)],
]:
    add_bullet(bullet)
add_tech_line("Python · TF-IDF · Logistic Regression · NLTK · Scikit-learn · Regex · Pandas · Geospatial Analysis")

# --- Project 6 ---
add_entry_head(
    "Zomato Dataset Analysis & Consumer Insights",
    "EDA & sentiment analysis",
    "Feb 2025",
)
for bullet in [
    [("Quantified a 65% online delivery adoption rate", True), (" and proved digital presence correlates with a 10–15% uplift in customer ratings.", False)],
    [("Segmented dining behaviour across customer cohorts", True), (" — surfaced that 55% of couples favour mid-range venues ($6–$18), a monetisable targeting signal.", False)],
    [("Ran sentiment analysis across thousands of reviews", True), (", identifying top 3 operational rating drivers (service speed, order accuracy, menu variety).", False)],
]:
    add_bullet(bullet)
add_tech_line("Python · Pandas · Matplotlib · Seaborn · Sentiment Analysis · EDA")


# ================== SKILLS ==================
add_heading_rule("Technical Skills")
skill_groups = [
    ("AI Agents", "LangChain · n8n · AI Agent Architecture · Prompt Engineering · Tool-First Routing · MCP"),
    ("LLMs & APIs", "GPT-4o · GPT-4.1-mini · Claude API · OpenAI Vision · OpenAI Images API · RESTful APIs · OAuth2"),
    ("Integrations", "Telegram Bot · LinkedIn REST · Gmail OAuth2 · Google Sheets/Docs · Supabase"),
    ("Machine Learning", "Python · Scikit-learn · XGBoost · Random Forest · NLP / TF-IDF · CNN / RNN"),
    ("Data & Cloud", "PostgreSQL · SQL · Power BI · Tableau · AWS · Google Cloud"),
    ("Tools", "Git / GitHub · VS Code · Jupyter · Docker (basics) · Linux"),
]
for label, val in skill_groups:
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(1)
    lab = sp.add_run(label + ": ")
    lab.bold = True
    lab.font.size = Pt(9.5)
    lab.font.color.rgb = INK
    v = sp.add_run(val)
    v.font.size = Pt(9.5)
    v.font.color.rgb = INK2


# ================== EDUCATION ==================
add_heading_rule("Education")
add_entry_head(
    "MSc Computer Science",
    "University of East London — London, UK",
    "",
)
ep = doc.add_paragraph()
ep.paragraph_format.space_before = Pt(0)
ep.paragraph_format.space_after = Pt(4)
e_run = ep.add_run("Focus: AI systems, automation, applied machine learning. Dissertation: ML-based GPS spoofing detection for drone delivery (99.95% accuracy).")
e_run.italic = True
e_run.font.size = Pt(9)
e_run.font.color.rgb = MUTED


# ================== CERTIFICATIONS ==================
add_heading_rule("Certifications")
certs = [
    "AWS Academy Cloud Foundations — Amazon Web Services",
    "Machine Learning, Deep Learning & Image Processing — MATLAB Onramps",
]
for c_text in certs:
    cp = doc.add_paragraph(style="List Bullet")
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(1)
    cp.paragraph_format.left_indent = Cm(0.5)
    cp.paragraph_format.first_line_indent = Cm(-0.3)
    cr = cp.add_run(c_text)
    cr.font.size = Pt(9.5)
    cr.font.color.rgb = INK2


# ================== SAVE ==================
doc.save(OUT)
print(f"Saved: {OUT}")
